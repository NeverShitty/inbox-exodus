"""
File Extractor module for Inbox Exodus
Specialized in extracting text from various file formats
"""
import os
import io
import email
import logging
import tempfile
import zipfile
import magic
import chardet
from typing import Dict, List, Tuple, Optional, Any, Union
from docx import Document
from pdfminer.high_level import extract_text as pdf_extract_text
from pdfminer.layout import LAParams

# Setup logger
logger = logging.getLogger(__name__)

class FileExtractor:
    """
    File content extractor for multiple document formats
    Supports PDF, DOCX, DOC, EML, TXT, and ZIP archives
    """
    
    def __init__(self, config=None):
        """
        Initialize the file extractor
        
        Args:
            config: Application configuration (optional)
        """
        from config import Config
        
        self.config = config or Config()
        self.temp_dir = self.config.temp_dir
        
        # Create temp directory if it doesn't exist
        os.makedirs(self.temp_dir, exist_ok=True)
    
    def extract_text(self, file_path: str, file_name: Optional[str] = None) -> Dict[str, Any]:
        """
        Extract text from a file based on its format
        
        Args:
            file_path: Path to the file
            file_name: Original name of the file (optional)
            
        Returns:
            Dict with extracted text and metadata
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return {"error": "File not found", "text": "", "metadata": {}}
            
        # Get file name if not provided
        if not file_name:
            file_name = os.path.basename(file_path)
            
        # Detect file type using python-magic
        file_type = self._detect_file_type(file_path)
        logger.info(f"Detected file type: {file_type} for {file_name}")
        
        # Extract text based on file type
        try:
            if "pdf" in file_type.lower():
                return self._extract_from_pdf(file_path, file_name)
            elif "word" in file_type.lower() or file_path.endswith(".docx"):
                return self._extract_from_docx(file_path, file_name)
            elif "email" in file_type.lower() or file_path.endswith(".eml"):
                return self._extract_from_email(file_path, file_name)
            elif "zip" in file_type.lower():
                return self._extract_from_zip(file_path, file_name)
            else:
                # Default to plain text extraction
                return self._extract_from_text(file_path, file_name)
        except Exception as e:
            logger.error(f"Error extracting text from {file_name}: {str(e)}")
            return {
                "error": f"Error extracting text: {str(e)}",
                "text": "",
                "metadata": {"file_name": file_name, "file_type": file_type}
            }
    
    def _detect_file_type(self, file_path: str) -> str:
        """
        Detect file type using python-magic
        
        Args:
            file_path: Path to the file
            
        Returns:
            String describing the file type
        """
        try:
            # Use python-magic to detect file type
            file_type = magic.from_file(file_path, mime=True)
            return file_type
        except Exception as e:
            logger.error(f"Error detecting file type: {str(e)}")
            # Try to guess from extension
            _, ext = os.path.splitext(file_path)
            if ext.lower() == '.pdf':
                return "application/pdf"
            elif ext.lower() in ['.docx', '.doc']:
                return "application/msword"
            elif ext.lower() == '.eml':
                return "message/rfc822"
            elif ext.lower() == '.zip':
                return "application/zip"
            else:
                return "text/plain"
    
    def _extract_from_pdf(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to the PDF file
            file_name: Name of the file
            
        Returns:
            Dict with extracted text and metadata
        """
        try:
            # Extract text from PDF
            text = pdf_extract_text(
                file_path,
                laparams=LAParams(),
                codec='utf-8'
            )
            
            metadata = {
                "file_name": file_name,
                "file_type": "PDF",
                "pages": 0  # We could count pages if needed
            }
            
            return {
                "text": text,
                "metadata": metadata
            }
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_name}: {str(e)}")
            return {
                "error": f"Error extracting text from PDF: {str(e)}",
                "text": "",
                "metadata": {"file_name": file_name, "file_type": "PDF"}
            }
    
    def _extract_from_docx(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to the DOCX file
            file_name: Name of the file
            
        Returns:
            Dict with extracted text and metadata
        """
        try:
            # Load the document
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extract properties
            properties = {
                "file_name": file_name,
                "file_type": "DOCX",
                "author": doc.core_properties.author or "Unknown",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "Unknown",
                "last_modified_by": doc.core_properties.last_modified_by or "Unknown",
                "paragraphs": len(doc.paragraphs)
            }
            
            return {
                "text": text,
                "metadata": properties
            }
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_name}: {str(e)}")
            return {
                "error": f"Error extracting text from DOCX: {str(e)}",
                "text": "",
                "metadata": {"file_name": file_name, "file_type": "DOCX"}
            }
    
    def _extract_from_email(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """
        Extract text from EML file
        
        Args:
            file_path: Path to the EML file
            file_name: Name of the file
            
        Returns:
            Dict with extracted text and metadata
        """
        try:
            # Parse email file
            with open(file_path, 'rb') as f:
                msg = email.message_from_binary_file(f)
            
            # Extract headers
            from_email = msg.get('From', 'Unknown')
            to_email = msg.get('To', 'Unknown')
            subject = msg.get('Subject', 'No Subject')
            date = msg.get('Date', 'Unknown')
            
            # Extract body
            body = ""
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    if content_type == "text/plain":
                        payload = part.get_payload(decode=True)
                        charset = part.get_content_charset()
                        if charset is None:
                            # Try to detect encoding
                            charset = chardet.detect(payload)['encoding']
                        if charset:
                            try:
                                body += payload.decode(charset, errors='replace')
                            except UnicodeDecodeError:
                                body += payload.decode('utf-8', errors='replace')
                        else:
                            body += payload.decode('utf-8', errors='replace')
            else:
                payload = msg.get_payload(decode=True)
                charset = msg.get_content_charset()
                if charset is None:
                    # Try to detect encoding
                    charset = chardet.detect(payload)['encoding']
                if charset:
                    try:
                        body = payload.decode(charset, errors='replace')
                    except UnicodeDecodeError:
                        body = payload.decode('utf-8', errors='replace')
                else:
                    body = payload.decode('utf-8', errors='replace')
            
            # Combine text for analysis
            text = f"Subject: {subject}\n\nFrom: {from_email}\nTo: {to_email}\nDate: {date}\n\n{body}"
            
            metadata = {
                "file_name": file_name,
                "file_type": "Email",
                "subject": subject,
                "from": from_email,
                "to": to_email,
                "date": date
            }
            
            return {
                "text": text,
                "metadata": metadata
            }
        except Exception as e:
            logger.error(f"Error extracting text from email {file_name}: {str(e)}")
            return {
                "error": f"Error extracting text from email: {str(e)}",
                "text": "",
                "metadata": {"file_name": file_name, "file_type": "Email"}
            }
    
    def _extract_from_zip(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """
        Extract files from ZIP archive and process each one
        
        Args:
            file_path: Path to the ZIP file
            file_name: Name of the ZIP file
            
        Returns:
            Dict with extracted text and metadata for all files
        """
        try:
            # Create temporary directory
            with tempfile.TemporaryDirectory(dir=self.temp_dir) as temp_dir:
                # Extract ZIP file
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)
                
                # Walk through extracted files
                files_content = []
                file_count = 0
                
                for root, _, files in os.walk(temp_dir):
                    for file in files:
                        if file.startswith('.') or file.startswith('__MACOSX'):
                            continue  # Skip hidden files and Mac OS X metadata
                            
                        file_count += 1
                        extracted_file_path = os.path.join(root, file)
                        
                        # Extract text from individual file
                        content = self.extract_text(extracted_file_path, file)
                        files_content.append(content)
                
                # Combine text from all files
                combined_text = ""
                for content in files_content:
                    if not content.get("error"):
                        file_metadata = content.get("metadata", {})
                        combined_text += f"\n\n--- File: {file_metadata.get('file_name', 'Unknown')} ---\n"
                        combined_text += content.get("text", "")
                
                metadata = {
                    "file_name": file_name,
                    "file_type": "ZIP Archive",
                    "file_count": file_count,
                    "files": [content.get("metadata", {}).get("file_name", "Unknown") for content in files_content],
                    "individual_results": files_content
                }
                
                return {
                    "text": combined_text,
                    "metadata": metadata
                }
        except Exception as e:
            logger.error(f"Error extracting text from ZIP {file_name}: {str(e)}")
            return {
                "error": f"Error extracting text from ZIP: {str(e)}",
                "text": "",
                "metadata": {"file_name": file_name, "file_type": "ZIP Archive"}
            }
    
    def _extract_from_text(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """
        Extract text from a plain text file with encoding detection
        
        Args:
            file_path: Path to the text file
            file_name: Name of the file
            
        Returns:
            Dict with extracted text and metadata
        """
        try:
            # Read file in binary mode
            with open(file_path, 'rb') as f:
                raw_data = f.read()
            
            # Detect encoding
            encoding_result = chardet.detect(raw_data)
            encoding = encoding_result['encoding'] or 'utf-8'
            confidence = encoding_result['confidence']
            
            # Decode text with detected encoding
            text = raw_data.decode(encoding, errors='replace')
            
            metadata = {
                "file_name": file_name,
                "file_type": "Text",
                "encoding": encoding,
                "encoding_confidence": confidence,
                "size": len(raw_data)
            }
            
            return {
                "text": text,
                "metadata": metadata
            }
        except Exception as e:
            logger.error(f"Error extracting text from {file_name}: {str(e)}")
            return {
                "error": f"Error extracting text: {str(e)}",
                "text": "",
                "metadata": {"file_name": file_name, "file_type": "Text"}
            }